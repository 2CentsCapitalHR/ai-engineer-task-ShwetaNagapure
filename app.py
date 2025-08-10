"""
RAG + Red-Flag Detection Flask app
- Upload .docx files
- RAG retrieval via Sentence-Transformers + FAISS over knowledgebase_docs.json
- LLM (Groq) for strict JSON output including red flags
- Annotate & highlight uploaded docs for red-flag locations
- Save annotated docs + JSON report for download
"""
import pdfkit  
import os
from datetime import datetime
import io
import re
import json
import uuid
import tempfile
from io import BytesIO
from typing import List, Tuple, Dict
from flask import Flask, request, render_template, redirect, url_for, send_file
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
from dotenv import load_dotenv
from groq import Groq

# Config + initialization

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise RuntimeError("Please set GROQ_API_KEY in .env")

GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-20b")
EMBED_MODEL_NAME = os.getenv("EMBED_MODEL", "all-MiniLM-L6-v2")

client = Groq(api_key=GROQ_API_KEY)
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100 MB limit for demo

# Simple in-memory job store (demo); use persistent storage in production
JOB_STORE: Dict[str, Dict] = {}


# Load knowledge + requirements

with open("knowledgebase_docs.json", "r", encoding="utf-8") as f:
    KNOWLEDGE_DOCS = json.load(f)  # list of {"title":..., "content":..."}

with open("requirement.json", "r", encoding="utf-8") as f:
    REQUIREMENTS = json.load(f)  # dict: process_name -> {description, required_documents}


# Embedding model + FAISS

embed_model = SentenceTransformer(EMBED_MODEL_NAME)

kb_texts = [d["content"] for d in KNOWLEDGE_DOCS]
kb_titles = [d["title"] for d in KNOWLEDGE_DOCS]
kb_embeddings = embed_model.encode(kb_texts, convert_to_numpy=True).astype("float32")

if kb_embeddings.size == 0:
    raise RuntimeError("Knowledge base embeddings are empty. Check knowledgebase_docs.json")

d = kb_embeddings.shape[1]
kb_index = faiss.IndexFlatL2(d)
kb_index.add(kb_embeddings)

# Build simple process embeddings to find candidate processes
process_names = list(REQUIREMENTS.keys())
process_texts = []
for p in process_names:
    text = REQUIREMENTS[p].get("description", "") + " " + " ".join(REQUIREMENTS[p].get("required_documents", []))
    process_texts.append(text)
process_embeddings = embed_model.encode(process_texts, convert_to_numpy=True).astype("float32")


# Helper functions

def extract_text_from_docx_fileobj(fileobj) -> str:
    """Return combined paragraphs' text from a docx file-like object."""
    doc = Document(fileobj)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip() != ""]
    return "\n".join(paras)

def retrieve_kb_docs(query_text: str, top_k: int = 3) -> List[Tuple[str,str]]:
    """Return top_k (title, content) results from KB using FAISS similarity."""
    q_emb = embed_model.encode([query_text], convert_to_numpy=True).astype("float32")
    D, I = kb_index.search(q_emb, top_k)
    res = []
    for idx in I[0]:
        if 0 <= idx < len(kb_texts):
            res.append((kb_titles[idx], kb_texts[idx]))
    return res

def identify_candidate_processes_and_matches(user_texts: List[str]) -> Tuple[List[str], Dict[str, List[str]]]:
    """
    Return candidate process names (top 3 similar) and a mapping process -> matched required docs (simple substring).
    """
    combined = " ".join(user_texts)
    q_emb = embed_model.encode([combined], convert_to_numpy=True).astype("float32")
    # compute L2 distances to process embeddings
    dists = np.linalg.norm(process_embeddings - q_emb, axis=1)
    top_idxs = list(np.argsort(dists)[:3])
    candidates = [process_names[i] for i in top_idxs]

    matches = {}
    combined_lower = combined.lower()
    for p in candidates:
        matches[p] = []
        for req in REQUIREMENTS[p].get("required_documents", []):
            if req.lower() in combined_lower:
                matches[p].append(req)
    return candidates, matches

def build_rag_redflag_prompt(user_docs_texts: List[str], retrieved_kb: List[Tuple[str,str]], candidate_processes: List[str], matched_docs: Dict[str,List[str]]) -> str:
    """
    Build a strict prompt instructing the LLM to:
    - Identify legal_process (choose from candidates if applicable)
    - List uploaded docs & missing docs (based on REQUIREMENTS)
    - Detect red flags (issues) with fields (document, section, issue, severity, suggestion, law_citation)
    Output MUST be a single JSON object matching the schema.
    """
    user_docs_combined = "\n\n".join(user_docs_texts)
    kb_context = "\n\n".join([f"Title: {t}\nContent: {c}" for t,c in retrieved_kb])
    matched_summary = "\n".join([f"{p}: {', '.join(v) if v else 'none detected'}" for p,v in matched_docs.items()])

    prompt = f"""
You are a corporate legal assistant. RESPOND ONLY with a single valid JSON object (no extra text) following this schema:

{{
  "legal_process": "<one of the candidate process names or free-text best fit>",
  "uploaded_documents": ["<list of uploaded filenames or doc types>"],
  "missing_documents": ["<list of required docs missing>"],
  "issues_found": [
    {{
      "document": "<filename or doc-type>",
      "section": "<short excerpt or heading where the issue appears>",
      "issue": "<description of the red flag>",
      "severity": "<low|medium|high>",
      "suggestion": "<short remediation suggestion>",
      "law_citation": "<optional law/regulation citation>"
    }}
  ],
  "summary_message": "<brief summary>"
}}

Candidate legal processes (top matches): {candidate_processes}

Quick matched required docs (from requirement.json): 
{matched_summary}

Relevant knowledgebase documents (use these to infer rules/clauses):
{kb_context}

User uploaded documents full text:
\"\"\"
{user_docs_combined}
\"\"\"

Tasks:
1) Decide the best-fit legal_process (prefer one of candidates if applicable).
2) Identify which required documents from requirement.json were uploaded (list them).
3) Identify missing required documents (list).
4) Thoroughly scan the user uploaded documents and detect any compliance "red flags" such as:
   - missing mandatory identifiers (IDs, registration numbers)
   - missing or unverified signatures
   - ambiguous or unconscionable clauses
   - jurisdiction or conflicting clause issues
   - missing required appendices or schedules
   For each red flag produce an entry in issues_found with document, a short section excerpt, a description, severity, suggestion and optional law citation.

Be concise in field values. DO NOT output anything other than the single JSON object. 
"""
    return prompt

def safe_parse_json_from_llm(llm_text: str) -> dict:
    """
    Robust attempt to parse JSON from LLM text.
    1) Try json.loads
    2) Extract first {...} block with regex and attempt to parse
    3) Try to fix trailing commas
    4) Return dict with _parse_error and raw if still failing
    """
    try:
        return json.loads(llm_text)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", llm_text, re.DOTALL)
        if not m:
            return {"_parse_error": "No JSON object found", "raw": llm_text}
        candidate = m.group(0)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            # try basic fixes
            cand_fixed = re.sub(r",\s*}", "}", candidate)
            cand_fixed = re.sub(r",\s*]", "]", cand_fixed)
            try:
                return json.loads(cand_fixed)
            except json.JSONDecodeError:
                return {"_parse_error": "Could not decode JSON after attempts", "raw_candidate": candidate, "raw_full": llm_text}

def highlight_paragraphs_and_insert_comments(doc: Document, issues: List[dict]):
    """
    For each issue:
      - find a paragraph that contains the issue['section'] text (case-insensitive)
      - highlight that paragraph's runs (yellow)
      - add an explanatory paragraph after it (visible comment)
    If section is empty or not found, append the comment at end.
    """
    for issue in issues:
        section = (issue.get("section") or "").strip()
        comment_text = f"RAG COMMENT - Issue: {issue.get('issue','')} | Severity: {issue.get('severity','')} | Suggestion: {issue.get('suggestion','')} | Law citation: {issue.get('law_citation','N/A')}"
        inserted = False
        if section:
            for p in doc.paragraphs:
                if section.lower() in p.text.lower():
                    # Highlight each run in this paragraph
                    for r in p.runs:
                        try:
                            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        except Exception:
                            pass
                    # Insert a new paragraph after this paragraph's block by appending at end (simpler, visible)
                    doc.add_paragraph(comment_text)
                    inserted = True
                    break
        if not inserted:
            # fallback: add at end
            doc.add_paragraph(comment_text)

# Flask routes

@app.route("/", methods=["GET", "POST"])
def upload_page():
    if request.method == "POST":
        files = request.files.getlist("documents")
        if not files or files[0].filename == "":
            return render_template("upload.html", error="Upload at least one .docx file.")

        # Read uploaded files: store original bytes, extracted text, Document objects
        docs_bytes = {}
        docs_texts = []
        docs_docobjs = {}
        filenames = []
        for f in files:
            fname = f.filename
            data = f.read()
            docs_bytes[fname] = data
            docs_texts.append(extract_text_from_docx_fileobj(BytesIO(data)))
            docs_docobjs[fname] = Document(BytesIO(data))
            filenames.append(fname)

        # Identify candidate processes and matched required docs (quick embedding-based)
        candidate_processes, matched_map = identify_candidate_processes_and_matches(docs_texts)

        # RAG retrieve top KB docs for context
        combined_text = " ".join(docs_texts)
        retrieved_kb = retrieve_kb_docs(combined_text, top_k=3)

        # Build strict prompt asking for red flags
        prompt = build_rag_redflag_prompt(docs_texts, retrieved_kb, candidate_processes, matched_map)

        messages = [
            {"role": "system", "content": "You are a corporate legal assistant. Reply ONLY with a single valid JSON object, nothing else."},
            {"role": "user", "content": prompt}
        ]

        # Call Groq LLM
        try:
            completion = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=messages,
                temperature=0.0,
                max_tokens=1800,
                top_p=1,
                stream=False
            )
        except Exception as e:
            return render_template("upload.html", error=f"LLM request failed: {e}")

        raw_llm_output = completion.choices[0].message.content

        # Try to parse JSON robustly
        rag_json = safe_parse_json_from_llm(raw_llm_output)
        if "_parse_error" in rag_json:
            # Show raw in template so user can debug
            return render_template("upload.html", error="Failed to parse LLM response as JSON.", raw_output=rag_json.get("raw_full") or rag_json.get("raw"))

        # At this point rag_json should have keys as per schema
        issues_found = rag_json.get("issues_found", []) if isinstance(rag_json, dict) else []

        # Create a new job to keep artifact paths
        job_id = str(uuid.uuid4())
        job_store_entry = {"annotated": {}, "highlighted": {}, "report": None}
        JOB_STORE[job_id] = job_store_entry

        # Annotate documents:
        # For each filename: collect issues that refer to it (or fallback logic)
        issues_by_file: Dict[str, List[dict]] = {}
        for issue in issues_found:
            docref = issue.get("document", "")
            placed = False
            # prefer exact filename
            if docref in docs_docobjs:
                issues_by_file.setdefault(docref, []).append(issue)
                placed = True
            else:
                # try partial match to filename
                for fname in docs_docobjs:
                    if docref and docref.lower() in fname.lower():
                        issues_by_file.setdefault(fname, []).append(issue)
                        placed = True
                        break
            if not placed:
                # if not matched, assign to first doc
                first = next(iter(docs_docobjs))
                issues_by_file.setdefault(first, []).append(issue)

        # For each doc: create two outputs:
        # 1) annotated copy (comments paragraphs + highlighted paragraphs)
        # 2) also save original copy (unchanged) if user wants original - we'll just save original bytes
        for fname, docobj in docs_docobjs.items():
            # work on a copy of the doc object to not mutate the object used elsewhere
            # python-docx Document object is mutable, but we have separate objects per filename already
            issues_for_file = issues_by_file.get(fname, [])

            # Annotated/highlighted doc (we'll highlight paragraphs and append visible comment paragraphs)
            doc_copy = docobj  # Document object already read from bytes; we'll modify and save as annotated

            if issues_for_file:
                highlight_paragraphs_and_insert_comments(doc_copy, issues_for_file)

            # Save annotated/highlighted doc to temp file
            annotated_name = f"annotated_{job_id}_{fname}"
            annotated_path = os.path.join(tempfile.gettempdir(), annotated_name)
            doc_copy.save(annotated_path)
            job_store_entry["annotated"][fname] = annotated_path

            # Save original uploaded bytes as a copy (so user can download original if needed)
            original_copy_path = os.path.join(tempfile.gettempdir(), f"original_{job_id}_{fname}")
            with open(original_copy_path, "wb") as orig_f:
                orig_f.write(docs_bytes[fname])
            job_store_entry["highlighted"].setdefault(fname, original_copy_path)

        # Save RAG JSON report
        report_path = os.path.join(tempfile.gettempdir(), f"rag_report_{job_id}.json")
        with open(report_path, "w", encoding="utf-8") as rf:
            json.dump(rag_json, rf, indent=2)
        job_store_entry["report"] = report_path

        # Redirect to results page
        return redirect(url_for("results_page", job_id=job_id))

    # show form
    return render_template("upload.html")

@app.route("/results/<job_id>")
def results_page(job_id):
    job = JOB_STORE.get(job_id)
    if not job:
        return "Job not found", 404
    # Load report for display
    with open(job["report"], "r", encoding="utf-8") as f:
        report = json.load(f)
    filenames = list(job["annotated"].keys())
    return render_template("results.html", report=report, filenames=filenames, job_id=job_id)

@app.route("/download_report/<job_id>")
def download_report(job_id):
    job = JOB_STORE.get(job_id)
    if not job:
        return "Job not found", 404
    path = job.get("report")
    if not path or not os.path.exists(path):
        return "Report not found", 404
    return send_file(path, as_attachment=True, download_name=f"rag_report_{job_id}.json", mimetype="application/json")

@app.route("/download_annotated/<job_id>/<filename>")
def download_annotated(job_id, filename):
    job = JOB_STORE.get(job_id)
    if not job:
        return "Job not found", 404
    annotated_map = job.get("annotated", {})
    path = annotated_map.get(filename)
    if not path or not os.path.exists(path):
        return "Annotated file not found", 404
    return send_file(path, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/download_original/<job_id>/<filename>")
def download_original(job_id, filename):
    job = JOB_STORE.get(job_id)
    if not job:
        return "Job not found", 404
    original_map = job.get("highlighted", {})
    path = original_map.get(filename)
    if not path or not os.path.exists(path):
        return "Original file not found", 404
    return send_file(path, as_attachment=True, download_name=f"original_{filename}", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True)
